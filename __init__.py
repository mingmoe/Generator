from . import shared

shared.inject_dependence()
from . import latex_helper

from typing import List
from datetime import datetime
import tempfile

from aqt import mw
from aqt.utils import showInfo, qconnect
from aqt.qt import *

from docx import Document
from docx.shared import Inches
from PIL import Image
from PIL.Image import Resampling
from munch import Munch

saying = "in a severe contest between intelligence, which presses forward, and an unworthy, timid ignorance obstructing our progress."


def generate_docx(files: List[str]) -> None:
    """
    弃用函数
    """
    config = mw.addonManager.getConfig(__name__)
    document = Document()
    document.add_heading(f"The Economist - {datetime.today().strftime('%Y %m %d')}", 0)
    document.add_paragraph(saying, style='Intense Quote')

    turn_page = False

    # 创建缩率图
    with tempfile.TemporaryDirectory() as tmpdirn:
        for file in files:
            with Image.open(file) as img:
                img.thumbnail((config['x'], config['y']), Resampling.HAMMING)
                img.save(os.path.join(tmpdirn, os.path.basename(file)) + ".png", "PNG")

            document.add_picture(os.path.join(tmpdirn, os.path.basename(file)) + ".png")

            if turn_page:
                turn_page = False
                document.add_page_break()
            else:
                document.add_paragraph("\n\n\n\n\n\n")
                turn_page = True

    # 设置页边距
    document.sections[0].left_margin = (Inches(0.25))
    document.sections[0].right_margin = (Inches(0.25))
    document.sections[0].top_margin = (Inches(0.25))
    document.sections[0].bottom_margin = (Inches(0.25))

    # 添加页眉
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    header = document.sections[0].header
    paragraph = header.paragraphs[0]

    # 设置文字
    run = paragraph.add_run()
    font = run.font
    font.italic = True
    font.size = Pt(12)
    paragraph.text = "MingMoe's Brainstorm.  Powered by MingMoe. With Love."
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存
    document.save((os.path.join(shared.root_dir, "output.docx")))
    os.startfile((os.path.join(shared.root_dir, "output.docx")))


def export_docx() -> None:
    """
    弃用函数
    """
    config = mw.addonManager.getConfig(__name__)

    found = mw.col.find_notes(f"{config['check_option']} tag:{config['tag']}")

    image_lists: List[str] = []

    for note_id in found:
        note = mw.col.get_note(note_id)
        files = mw.col.media.files_in_str(note.mid, note[config['field']])
        for file in files:
            file = mw.col.media.escape_media_filenames(file, True)
            image_lists.append(os.path.join(mw.col.media.dir(), file))

    showInfo(f"get cards: {len(found)}")
    generate_docx(image_lists)


thum_count: int = 0


def create_thum_for_file(file: str, size: (int, int)) -> str:
    """
    为文件创建缩率图
    :param file:
    :param size:
    :return: 路径
    """
    # 创建缩率图
    global thum_count
    thum_count += 1
    n = os.path.join(shared.output_dir, str(thum_count)) + ".png"
    with Image.open(file) as img:
        img.thumbnail(size, Resampling.HAMMING)
        img.save(n, "PNG")
    return n


import imgkit

def create_pdf_for_html(html: str) -> str:
    global thum_count
    thum_count += 1
    n = os.path.join(shared.output_dir, str(thum_count)) + ".png"

    options = {
        'format':"png",
    }

    generated = shared.merge_html(html)

    with open(os.path.join(shared.output_dir,"generated.html"),encoding="utf-8",mode="w") as f:
        f.truncate(0)
        f.write(generated)

    imgkit.from_string(generated, n, options)
    return n


def generate_latex(items: List[shared.LaTeXItem]) -> None:
    latex = latex_helper.LaTeXDocument()

    latex.debug = config = mw.addonManager.getConfig(__name__)['debug']

    path = latex_helper.str_to_latex_str(shared.output_dir.replace('\\', '/'))
    latex.first_source = f"\\graphicspath{{{path}}}"

    for item in items:
        code = f"\\includegraphics{{{os.path.basename(item.image_path)}}}\n"

        if item.pdf_path is not None:
            code += f"\\includegraphics{{{os.path.basename(item.pdf_path)}}}\n"

        latex.pages.append(code)

    opts = latex.build()

    if opts is not None:
        showInfo(opts)
    try:
        os.startfile(latex.get_output_pdf_path())
    except BaseException:
        pass


def export_pdf() -> None:
    global thum_count
    # clear the counter
    thum_count = 0

    # clear the output directory
    shared.clear_output_dir()

    config_dict: dict = mw.addonManager.getConfig(__name__)

    config: shared.Config = Munch(config_dict)

    found_list: List[shared.LaTeXItem] = []

    for deck in config.decks:
        deck: shared.DeckConfig = Munch(deck)

        found = mw.col.find_notes(f"deck:{deck.name} {deck.search_option}")

        for note_id in found:
            try:
                note = mw.col.get_note(note_id)
                files = mw.col.media.files_in_str(note.mid, note[deck.found_image_at_field])

                note_text = note[deck.found_string_at_field]
#                if len(note_text) != 0:
#                    text_path = create_pdf_for_html(note_text)
                #else:
                    #text_path = None
                note_text=None

                # translate file
                for file in files:
                    image_path = \
                        create_thum_for_file(os.path.join(mw.col.media.dir(), file),
                                             (deck.image_x, deck.image_y))

                    latex_item: shared.LaTeXItem = shared.LaTeXItem()
                    # latex_item.pdf_path = text_path
                    latex_item.pdf_path = None
                    latex_item.image_path = image_path
                    found_list.append(latex_item)
            except BaseException as ex:
                showInfo(f"found error at card " + str(note_id))
                showInfo(str(ex))
                showInfo(''.join(traceback.TracebackException.from_exception(ex).format()))
                break

    showInfo(f"get cards: {len(found_list)}")
    generate_latex(found_list)


action = QAction("export cards to docx", mw)

qconnect(action.triggered, export_pdf)

mw.form.menuTools.addAction(action)
